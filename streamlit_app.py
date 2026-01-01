import streamlit as st
import google.generativeai as genai
import PyPDF2
from docx import Document
from io import BytesIO

# Configuración de la página
st.set_page_config(page_title="Asistente de Competencias", layout="wide")

st.title("🛠️ Sistema de Elaboración y Revisión de Estándares")

# Sidebar
with st.sidebar:
    st.header("1. Configuración")
    api_key = st.text_input("Pega tu API Key de Gemini:", type="password")
    modo = st.radio("2. Elige función:", ["Revisión de Documentos", "Creación desde Cero"])
    st.divider()
    st.caption("Asegúrate de haber creado tu API Key en Google AI Studio.")

# Validación de API Key
if not api_key:
    st.info("👋 ¡Hola! Para empezar, pega tu API Key en la barra de la izquierda.")
    st.stop()

# Configuración del modelo
try:
    genai.configure(api_key=api_key)
    # Usamos gemini-1.5-flash que es el más estable actualmente
    model = genai.GenerativeModel('gemini-1.5-flash')
except Exception as e:
    st.error(f"Error al configurar la IA: {e}")
    st.stop()

def extraer_texto_pdf(archivo_pdf):
    try:
        lector = PyPDF2.PdfReader(archivo_pdf)
        texto = ""
        for pagina in lector.pages:
            texto += pagina.extract_text()
        return texto
    except Exception as e:
        st.error(f"No pude leer el PDF: {e}")
        return None

def crear_word(contenido):
    doc = Document()
    doc.add_heading('Resultado del Análisis de Competencias', 0)
    doc.add_paragraph(contenido)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- CUERPO DE LA APP ---
archivo_estandar = st.file_uploader("📂 Paso 1: Sube el Estándar de Competencia (PDF)", type="pdf")

if archivo_estandar:
    texto_estandar = extraer_texto_pdf(archivo_estandar)
    
    if texto_estandar:
        st.success("✅ Estándar cargado correctamente.")
        
        if modo == "Revisión de Documentos":
            archivo_usuario = st.file_uploader("📄 Paso 2: Sube el Documento a Revisar (PDF)", type="pdf")
            if archivo_usuario:
                if st.button("🔍 Iniciar Auditoría"):
                    with st.spinner("Analizando..."):
                        try:
                            prompt = f"Actúa como auditor. Compara este documento con este estándar: {texto_estandar[:5000]}. Genera una tabla de cumplimiento."
                            # Agregamos manejo de error aquí
                            response = model.generate_content(prompt)
                            st.markdown(response.text)
                        except Exception as e:
                            st.error(f"Hubo un error con la IA: {e}. Intenta revisar si tu API Key es correcta.")

        elif modo == "Creación desde Cero":
            if st.button("📝 Generar Preguntas de Diagnóstico"):
                with st.spinner("Generando entrevista..."):
                    try:
                        # Limitamos el texto para evitar errores de capacidad
                        prompt_preguntas = f"Basado en este estándar de competencia: {texto_estandar[:4000]}, genera 5 preguntas clave para obtener información de los productos y desempeños que pide el estándar."
                        response = model.generate_content(prompt_preguntas)
                        st.session_state['preguntas'] = response.text
                    except Exception as e:
                        st.error(f"Error al generar preguntas: {e}")

            if 'preguntas' in st.session_state:
                st.markdown("### Responde estas preguntas:")
                st.info(st.session_state['preguntas'])
                respuestas = st.text_area("Escribe aquí tus respuestas:")
                
                if st.button("✨ Redactar Documento Final"):
                    with st.spinner("Redactando tabla técnica..."):
                        try:
                            prompt_final = f"Con el estándar {texto_estandar[:2000]} y estas respuestas: {respuestas}, crea una tabla técnica profesional con Productos, Desempeños y Actitudes/Valores."
                            resultado = model.generate_content(prompt_final)
                            st.markdown(resultado.text)
                            st.download_button("Descargar en Word", crear_word(resultado.text), "Producto.docx")
                        except Exception as e:
                            st.error(f"Error al redactar: {e}")
else:
    st.info("Para comenzar, por favor sube el archivo PDF del estándar que vamos a usar como base.")
